import io
import re
from sqlalchemy.orm import Session
from sqlalchemy import select, func, case, exists
from docx import Document
from database.db import IS_SQLITE
from models.soporte import SoporteDocumento, SoporteChunk
from schemas.soporte import SoporteInfoResponse
from services import ai_provider

MAX_DOCS_CONTEXTO = 2
MAX_CHARS_POR_DOC = 4000
MAX_PALABRAS_CLAVE = 12

TAM_CHUNK = 1500
SOLAPE_CHUNK = 200
MAX_CHUNKS_POR_DOC = 400  # acota el costo de embeddings para documentos enormes
MAX_CHUNKS_CONTEXTO = 6

_BLOQUE_CONTEXTO_TEMPLATE = """

## Documentos de soporte cargados por la entidad
A continuación hay extractos de documentos internos que la entidad cargó como referencia. Cítalos por su nombre de archivo cuando los uses. Si alguno contiene el texto literal de un artículo o norma, puedes citarlo entre comillas tal cual aparece; si no, no inventes una redacción exacta. Si no aportan nada a la pregunta, ignóralos.

{contexto}"""

_STOPWORDS = {
    "de", "la", "el", "en", "y", "a", "que", "los", "las", "un", "una", "es",
    "del", "se", "por", "con", "para", "su", "al", "lo", "como", "más", "o",
    "pero", "sus", "le", "ya", "este", "esta", "entre", "sin", "sobre", "ser",
    "qué", "cómo", "cuál", "cuáles", "cuándo", "dónde", "puede", "debe",
}


def _extraer_texto(file_bytes: bytes, tipo_archivo: str) -> str:
    if tipo_archivo == "pdf":
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    if tipo_archivo == "docx":
        doc = Document(io.BytesIO(file_bytes))
        partes = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    partes.append(cell.text)
        return "\n".join(partes)
    return file_bytes.decode("utf-8", errors="ignore")


def extraer_texto_archivo(file_bytes: bytes, filename: str) -> str:
    """Extrae el texto de un PDF/DOCX/TXT sin persistirlo. Lo usan flujos puntuales
    como el informe técnico de cuarta clase, donde el texto se muestra al usuario
    para que elija los párrafos que va a incorporar a la motivada."""
    tipo_archivo = filename.rsplit(".", 1)[-1].lower()
    return _extraer_texto(file_bytes, tipo_archivo)


def _a_response(doc: SoporteDocumento) -> SoporteInfoResponse:
    return SoporteInfoResponse(
        id=doc.id,
        nombre_original=doc.nombre_original,
        tipo_archivo=doc.tipo_archivo,
        tamano_bytes=doc.tamano_bytes,
        longitud_texto=len(doc.contenido_texto),
        fecha_subida=doc.fecha_subida,
    )


def _chunkear_texto(texto: str) -> list[str]:
    texto = texto.strip()
    if not texto:
        return []
    chunks = []
    inicio = 0
    n = len(texto)
    while inicio < n and len(chunks) < MAX_CHUNKS_POR_DOC:
        fin = min(inicio + TAM_CHUNK, n)
        chunks.append(texto[inicio:fin])
        if fin == n:
            break
        inicio = fin - SOLAPE_CHUNK
    return chunks


def _indexar_chunks(db: Session, doc: SoporteDocumento) -> None:
    """Genera embeddings para los fragmentos del documento y los guarda para
    busqueda semantica (RAG). Falla en silencio: si los embeddings no estan
    disponibles, el documento sigue accesible por la busqueda por palabras
    clave (_buscar_contexto_legacy)."""
    fragmentos = _chunkear_texto(doc.contenido_texto)
    if not fragmentos:
        return
    embeddings = ai_provider.embed_texts(fragmentos, task_type="retrieval_document")
    for orden, (texto, vector) in enumerate(zip(fragmentos, embeddings)):
        db.add(SoporteChunk(soporte_id=doc.id, orden=orden, texto=texto, embedding=vector))
    db.commit()


def guardar_soporte(db: Session, file_bytes: bytes, filename: str) -> SoporteInfoResponse:
    tipo_archivo = filename.rsplit(".", 1)[-1].lower()
    texto = _extraer_texto(file_bytes, tipo_archivo)
    if not texto.strip():
        raise ValueError("No se pudo extraer texto del documento")

    doc = SoporteDocumento(
        nombre_original=filename,
        tipo_archivo=tipo_archivo,
        contenido_texto=texto,
        tamano_bytes=len(file_bytes),
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)

    if not IS_SQLITE:
        try:
            _indexar_chunks(db, doc)
        except Exception:
            pass

    return _a_response(doc)


def listar_soportes(db: Session) -> list[SoporteInfoResponse]:
    filas = db.execute(
        select(
            SoporteDocumento.id,
            SoporteDocumento.nombre_original,
            SoporteDocumento.tipo_archivo,
            SoporteDocumento.tamano_bytes,
            func.length(SoporteDocumento.contenido_texto).label("longitud_texto"),
            SoporteDocumento.fecha_subida,
        ).order_by(SoporteDocumento.fecha_subida.desc())
    ).all()
    return [
        SoporteInfoResponse(
            id=f.id,
            nombre_original=f.nombre_original,
            tipo_archivo=f.tipo_archivo,
            tamano_bytes=f.tamano_bytes,
            longitud_texto=f.longitud_texto,
            fecha_subida=f.fecha_subida,
        )
        for f in filas
    ]


def eliminar_soporte(db: Session, soporte_id: int) -> bool:
    doc = db.get(SoporteDocumento, soporte_id)
    if not doc:
        return False
    db.delete(doc)  # los chunks se eliminan en cascada (ON DELETE CASCADE)
    db.commit()
    return True


def _palabras_clave(texto: str) -> list[str]:
    palabras = re.findall(r"[a-záéíóúñ0-9]+", texto.lower())
    return [p for p in palabras if p not in _STOPWORDS and len(p) > 2]


def _buscar_contexto_legacy(db: Session, pregunta: str) -> str:
    """Busqueda por solapamiento de palabras clave (sin embeddings). Se usa
    como respaldo si los embeddings fallan o en SQLite (desarrollo local,
    donde no esta disponible la extension pgvector)."""
    claves = list(set(_palabras_clave(pregunta)))[:MAX_PALABRAS_CLAVE]
    if not claves:
        return ""

    contenido_lower = func.lower(SoporteDocumento.contenido_texto)
    puntaje = sum(
        case((contenido_lower.like(f"%{palabra}%"), 1), else_=0) for palabra in claves
    )

    subq = select(
        SoporteDocumento.nombre_original.label("nombre"),
        func.substr(SoporteDocumento.contenido_texto, 1, MAX_CHARS_POR_DOC).label("extracto"),
        puntaje.label("puntaje"),
    ).subquery()

    filas = db.execute(
        select(subq.c.nombre, subq.c.extracto)
        .where(subq.c.puntaje > 0)
        .order_by(subq.c.puntaje.desc())
        .limit(MAX_DOCS_CONTEXTO)
    ).all()

    if not filas:
        return ""

    return "\n\n".join(f"--- {nombre} ---\n{extracto}" for nombre, extracto in filas)


def _backfill_chunks_faltantes(db: Session) -> None:
    """Indexa documentos subidos antes de activar el RAG, o que fallaron al
    indexarse en su momento (p. ej. por una caida temporal de la API de
    embeddings)."""
    pendientes = db.scalars(
        select(SoporteDocumento).where(
            ~exists().where(SoporteChunk.soporte_id == SoporteDocumento.id)
        )
    ).all()
    for doc in pendientes:
        try:
            _indexar_chunks(db, doc)
        except Exception:
            # Una falla aqui (p. ej. la API de embeddings) deja la sesion en
            # estado "pending rollback" — sin esto, la siguiente consulta en
            # esta misma sesion (la busqueda RAG, o incluso su fallback por
            # palabras clave) tambien fallaria.
            db.rollback()


def buscar_contexto_relevante(db: Session, pregunta: str) -> str:
    """Busqueda semantica (RAG): embebe la pregunta y recupera, por similitud
    de coseno calculada en la base de datos (pgvector), los fragmentos de
    los documentos de soporte mas relevantes — sin transferir ni procesar en
    Python el contenido completo de cada documento."""
    if IS_SQLITE:
        return _buscar_contexto_legacy(db, pregunta)

    try:
        _backfill_chunks_faltantes(db)
        vector_pregunta = ai_provider.embed_texts([pregunta], task_type="retrieval_query")[0]

        filas = db.execute(
            select(SoporteDocumento.nombre_original, SoporteChunk.texto)
            .join(SoporteDocumento, SoporteChunk.soporte_id == SoporteDocumento.id)
            .order_by(SoporteChunk.embedding.cosine_distance(vector_pregunta))
            .limit(MAX_CHUNKS_CONTEXTO)
        ).all()
    except Exception:
        return _buscar_contexto_legacy(db, pregunta)

    if not filas:
        return ""

    return "\n\n".join(f"--- {nombre} ---\n{texto}" for nombre, texto in filas)


def construir_bloque_contexto(contexto: str) -> str:
    """Envuelve un extracto de soportes con las instrucciones de uso, listo para anexar a un system prompt."""
    if not contexto:
        return ""
    return _BLOQUE_CONTEXTO_TEMPLATE.format(contexto=contexto)
