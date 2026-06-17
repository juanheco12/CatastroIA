import re
from sqlalchemy.orm import Session
from sqlalchemy import select
from docx import Document
from models.template import TemplateActivo


def save_template(db: Session, file_bytes: bytes, original_name: str) -> dict:
    """Persists an uploaded .docx template in the database and returns its metadata."""
    campos = detect_template_fields(file_bytes)

    db.query(TemplateActivo).delete()
    plantilla = TemplateActivo(
        nombre_original=original_name,
        contenido=file_bytes,
        tamano_bytes=len(file_bytes),
        campos_detectados=",".join(campos),
    )
    db.add(plantilla)
    db.commit()
    db.refresh(plantilla)

    return {
        "nombre_original": plantilla.nombre_original,
        "fecha_subida": plantilla.fecha_subida.isoformat(),
        "campos_detectados": campos,
        "tamano_bytes": plantilla.tamano_bytes,
    }


def detect_template_fields(file_bytes: bytes) -> list[str]:
    """Extracts {{FIELD}} placeholders from a .docx file."""
    import io
    doc = Document(io.BytesIO(file_bytes))
    full_text = " ".join(p.text for p in doc.paragraphs)
    # Also scan tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += " " + cell.text
    return list(set(re.findall(r"\{\{(\w+)\}\}", full_text)))


def get_template_info(db: Session) -> dict:
    """Returns metadata about the currently stored template."""
    plantilla = db.scalars(select(TemplateActivo)).first()
    if not plantilla:
        return {"existe": False, "nombre": "", "campos_detectados": [], "tamano_bytes": 0}

    return {
        "existe": True,
        "nombre": plantilla.nombre_original,
        "campos_detectados": plantilla.campos_detectados.split(",") if plantilla.campos_detectados else [],
        "tamano_bytes": plantilla.tamano_bytes,
        "fecha_subida": plantilla.fecha_subida.isoformat(),
    }


def load_template_bytes(db: Session) -> bytes | None:
    plantilla = db.scalars(select(TemplateActivo)).first()
    return plantilla.contenido if plantilla else None
