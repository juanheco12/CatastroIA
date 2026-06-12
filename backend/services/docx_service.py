import io
import re
from datetime import date, datetime
from docx import Document
from docx.oxml.ns import qn
from copy import deepcopy
from schemas.tercera_clase import TerceraClaseInput
from services.template_service import load_template_bytes


# Default template built in memory if no custom template is uploaded
DEFAULT_TEMPLATE_TEXT = """RESOLUCIÓN N° {{NUMERO_EXPEDIENTE}}

Por la cual se decide la solicitud de mutación catastral de Tercera Clase
(Incorporación de Construcción) sobre el predio {{NUMERO_PREDIO}}

EL RESPONSABLE DE CATASTRO

En ejercicio de sus atribuciones legales y en especial las conferidas por la
Ley 388 de 1997, el Decreto 1420 de 1998 y la Resolución 1040 de 2009 del IGAC,

CONSIDERANDO

{{MOTIVADA}}

RESUELVE

ARTÍCULO PRIMERO: Incorporar al inventario catastral la construcción ubicada en
{{DIRECCION_CONSTRUCCION}}, municipio de {{MUNICIPIO}}, con un área construida de
{{AREA_CONSTRUIDA}} m², de propiedad de {{PROPIETARIO_NOMBRE}}, identificado(a)
con {{TIPO_DOCUMENTO}} N° {{NUMERO_DOCUMENTO}}.

ARTÍCULO SEGUNDO: La presente resolución rige a partir de su ejecutoria.

NOTIFÍQUESE Y CÚMPLASE

Dada en {{MUNICIPIO}}, a los {{DIA_RESOLUCION}} días del mes de {{MES_RESOLUCION}}
del año {{ANIO_RESOLUCION}}.

{{INSPECTOR_RESPONSABLE}}
{{CARGO_INSPECTOR}}
"""

MESES = [
    "", "enero", "febrero", "marzo", "abril", "mayo", "junio",
    "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
]


def _build_replacements(data: TerceraClaseInput, texto_motivada: str) -> dict[str, str]:
    today = date.today()
    return {
        "NUMERO_EXPEDIENTE": data.numero_expediente,
        "NUMERO_PREDIO": data.numero_predio,
        "MATRICULA_INMOBILIARIA": data.matricula_inmobiliaria or "N/A",
        "PROPIETARIO_NOMBRE": data.propietario.nombre_completo,
        "TIPO_DOCUMENTO": data.propietario.tipo_documento,
        "NUMERO_DOCUMENTO": data.propietario.numero_documento,
        "DIRECCION_CONSTRUCCION": data.construccion.direccion,
        "MUNICIPIO": data.construccion.municipio,
        "DEPARTAMENTO": data.construccion.departamento,
        "AREA_CONSTRUIDA": str(data.construccion.area_construida_m2),
        "ANIO_CONSTRUCCION": str(data.construccion.anio_construccion),
        "NUMERO_PISOS": str(data.construccion.numero_pisos),
        "MATERIALES": data.construccion.materiales_predominantes,
        "USO_CONSTRUCCION": data.construccion.uso_construccion,
        "DESTINO_ECONOMICO": data.construccion.destino_economico,
        "FECHA_SOLICITUD": str(data.fecha_solicitud),
        "INSPECTOR_RESPONSABLE": data.inspector_responsable,
        "CARGO_INSPECTOR": data.cargo_inspector,
        "MOTIVADA": texto_motivada,
        "DIA_RESOLUCION": str(today.day),
        "MES_RESOLUCION": MESES[today.month],
        "ANIO_RESOLUCION": str(today.year),
    }


def _replace_in_paragraph(paragraph, replacements: dict[str, str]) -> None:
    """Replaces {{KEY}} tokens preserving paragraph formatting as much as possible."""
    full_text = "".join(run.text for run in paragraph.runs)
    if "{{" not in full_text:
        return

    new_text = full_text
    for key, value in replacements.items():
        new_text = new_text.replace(f"{{{{{key}}}}}", value)

    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""


def generate_docx(data: TerceraClaseInput, texto_motivada: str) -> bytes:
    """
    Injects form data + generated motivada into the Word template.
    Falls back to a built-in default template when no custom one is uploaded.
    """
    replacements = _build_replacements(data, texto_motivada)

    template_bytes = load_template_bytes()

    if template_bytes:
        doc = Document(io.BytesIO(template_bytes))
        for paragraph in doc.paragraphs:
            _replace_in_paragraph(paragraph, replacements)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        _replace_in_paragraph(paragraph, replacements)
    else:
        # Build a clean document from the default template text
        doc = Document()
        doc.core_properties.author = "CatIA - Sistema Catastral"
        doc.core_properties.title = f"Resolución {data.numero_expediente}"

        filled_text = DEFAULT_TEMPLATE_TEXT
        for key, value in replacements.items():
            filled_text = filled_text.replace(f"{{{{{key}}}}}", value)

        for line in filled_text.strip().split("\n"):
            doc.add_paragraph(line)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
