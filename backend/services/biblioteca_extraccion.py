"""Extraccion de texto plano y parrafos de un .docx para la biblioteca de
motivadas. El texto plano es el mismo en ingestion y en generacion (metodo
determinista y puro) — los offsets de CampoVariablePlantilla viven sobre
este texto, nunca sobre el XML."""
import io
from docx import Document
from docx.text.paragraph import Paragraph

PARRAFO_SEPARATOR = "\n"


def _parrafos_en_orden(doc: Document) -> list[Paragraph]:
    """Parrafos del cuerpo + celdas de tabla, en orden de documento — mismo
    orden usado tanto para extraer texto como para mapear runs."""
    parrafos = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parrafos.extend(cell.paragraphs)
    return parrafos


def extraer_texto_plano(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    parrafos = _parrafos_en_orden(doc)
    return PARRAFO_SEPARATOR.join(p.text for p in parrafos)


def extraer_parrafos_con_runs(file_bytes: bytes) -> tuple[Document, list[Paragraph]]:
    """Devuelve el Document abierto (necesario para luego doc.save()) y la
    lista de Paragraph en el mismo orden que extraer_texto_plano, para que
    el motor de reemplazo pueda mapear offsets de texto plano a runs reales."""
    doc = Document(io.BytesIO(file_bytes))
    return doc, _parrafos_en_orden(doc)